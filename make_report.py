from docx import Document
from docx.shared import Inches
import os

doc = Document()
doc.add_heading('Automobile Service Project - Test Report', level=1)
doc.add_paragraph('Server logs and screenshots attached below (if you add files into the project directory).')

if os.path.exists('screenshot_docs.png'):
    doc.add_heading('API Docs', level=2)
    doc.add_picture('screenshot_docs.png', width=Inches(6))

if os.path.exists('pytest_output.png'):
    doc.add_heading('Test Results', level=2)
    doc.add_picture('pytest_output.png', width=Inches(6))

doc.add_heading('Notes', level=2)
doc.add_paragraph('Project completed and validated per business rules.')

doc.save('Test_Report.docx')
print('Saved Test_Report.docx')
